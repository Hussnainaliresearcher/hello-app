import streamlit as st
import os
from PyPDF2 import PdfReader
import docx
from langchain_community.chat_models import ChatOpenAI
from langchain.vectorstores import Qdrant
import random
from datetime import datetime
from langchain import PromptTemplate
from langchain.chains import RetrievalQA
import string
from qdrant_client import QdrantClient
from dotenv import load_dotenv
from langchain.embeddings import HuggingFaceEmbeddings
from langchain.text_splitter import CharacterTextSplitter
from streamlit_chat import message
from langchain.callbacks import get_openai_callback
from langchain.docstore.document import Document

# Load environment variables
load_dotenv()

# Load API keys from Streamlit secrets
openapi_key = st.secrets["OPENAI_API_KEY"]
qdrant_url = st.secrets["QDRANT_URL"]
qdrant_api_key = st.secrets["QDRANT_API_KEY"]

# Define embeddings model
embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")


def main():
    st.set_page_config(page_title="Q/A with your file")
    st.header("Pakistan Organic Association")

    if "conversation" not in st.session_state:
        st.session_state.conversation = None
    if "chat_history" not in st.session_state:
        st.session_state.chat_history = []
    if "processComplete" not in st.session_state:
        st.session_state.processComplete = None

    with st.sidebar:
        uploaded_files = st.file_uploader("Upload your file", type=['pdf'], accept_multiple_files=True)
        openai_api_key = openapi_key
        process = st.button("Process")

    if process:
        with st.spinner('Processing...'):
            if not openai_api_key:
                st.info("Please add your OpenAI API key to continue.")
                st.stop()

            text_chunks_list = []
            for uploaded_file in uploaded_files:
                file_name = uploaded_file.name
                file_text = get_files_text(uploaded_file)
                text_chunks = get_text_chunks(file_text, file_name)
                text_chunks_list.extend(text_chunks)

            # Create vector store
            curr_date = str(datetime.now())
            collection_name = "".join(random.choices(string.ascii_letters, k=4)) + curr_date.split('.')[0].replace(':', '-').replace(" ", 'T')
            vectorstore = get_vectorstore(text_chunks_list, collection_name)

            st.write("Vector Store Created...")

            # Create QA chain
            num_chunks = 4
            st.session_state.conversation = get_qa_chain(vectorstore, num_chunks)
            st.session_state.processComplete = True

    if st.session_state.processComplete:
        user_question = st.chat_input("Ask a question about your files.")
        if user_question:
            handle_user_input(user_question)


# Function to get text from uploaded file
def get_files_text(uploaded_file):
    text = ""
    file_extension = os.path.splitext(uploaded_file.name)[1]

    if file_extension == ".pdf":
        text += get_pdf_text(uploaded_file)
    elif file_extension == ".docx":
        text += get_docx_text(uploaded_file)

    return text


# Function to extract text from PDF
def get_pdf_text(pdf):
    pdf_reader = PdfReader(pdf)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text


# Function to extract text from DOCX
def get_docx_text(file):
    doc = docx.Document(file)
    return " ".join([para.text for para in doc.paragraphs])


# Function to split text into chunks
def get_text_chunks(text, filename):
    text_splitter = CharacterTextSplitter(
        separator="\n",
        chunk_size=80,
        chunk_overlap=20,
        length_function=len
    )
    chunks = text_splitter.split_text(text)
    return [Document(page_content=chunk, metadata={"source": filename}) for chunk in chunks]


# Function to create vector store
def get_vectorstore(text_chunks, collection_name):
    try:
        knowledge_base = Qdrant.from_documents(
            documents=text_chunks,
            embedding=embeddings,
            url=qdrant_url,
            prefer_grpc=True,
            api_key=qdrant_api_key,
            collection_name=collection_name,
        )
        return knowledge_base
    except Exception as e:
        st.write(f"Error: {e}")
        return None


# Function to create QA chain
def get_qa_chain(vectorstore, num_chunks):
    if vectorstore is None:
        st.error("Vector store creation failed.")
        return None

    qa = RetrievalQA.from_chain_type(
        llm=ChatOpenAI(model="gpt-3.5-turbo"),
        chain_type="stuff",
        retriever=vectorstore.as_retriever(search_type="similarity", search_kwargs={"k": num_chunks}),
        return_source_documents=True
    )
    return qa


# Function to handle user input
def handle_user_input(user_question):
    with st.spinner('Generating response...'):
        result = st.session_state.conversation({"query": user_question})
        response = result.get('result', "No response generated.")
        source = result['source_documents'][0].metadata.get('source', "Unknown Source")

    st.session_state.chat_history.append(user_question)
    st.session_state.chat_history.append(f"{response} \n Source Document: {source}")

    response_container = st.container()

    with response_container:
        for i, message_text in enumerate(st.session_state.chat_history):
            if i % 2 == 0:
                message(message_text, is_user=True, key=str(i))
            else:
                message(message_text, key=str(i))


if __name__ == '__main__':
    main()
